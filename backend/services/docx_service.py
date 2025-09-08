from docx import Document
from typing import Union, Dict

def save_resume_docx(data: Union[str, Dict], filename="resume.docx"):
    doc = Document()

    if isinstance(data, str):
        # Fallback: add raw text to the doc
        for line in data.splitlines():
            doc.add_paragraph(line)
        doc.save(filename)
        return filename

    # Structured data path
    name = data.get('name', '')
    if name:
        doc.add_heading(name, 0)

    contact_parts = []
    if data.get('phone'): contact_parts.append(data.get('phone'))
    if data.get('email'): contact_parts.append(data.get('email'))
    if data.get('linkedin'): contact_parts.append(data.get('linkedin'))
    if contact_parts:
        doc.add_paragraph(' | '.join(contact_parts))

    if data.get('summary'):
        doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        for line in data.get('summary').splitlines():
            doc.add_paragraph(line)

    if data.get('education'):
        doc.add_heading('EDUCATION', level=1)
        doc.add_paragraph(data.get('education'))

    if data.get('skills'):
        doc.add_heading('SKILLS', level=1)
        skills = data.get('skills')
        if isinstance(skills, (list, tuple)):
            skills = ', '.join(skills)
        doc.add_paragraph(skills)

    if data.get('experience'):
        doc.add_heading('WORK EXPERIENCE', level=1)
        for line in data.get('experience').splitlines():
            doc.add_paragraph(line)

    if data.get('projects'):
        doc.add_heading('PROJECTS', level=1)
        doc.add_paragraph(data.get('projects'))

    if data.get('certifications'):
        doc.add_heading('CERTIFICATIONS', level=1)
        doc.add_paragraph(data.get('certifications'))

    # Fallback: include generated_text if provided
    if data.get('generated_text') and not any([data.get(k) for k in ['summary','experience','education','projects','certifications']]):
        doc.add_heading('GENERATED RESUME', level=1)
        for line in data.get('generated_text').splitlines():
            doc.add_paragraph(line)

    doc.save(filename)
    return filename
